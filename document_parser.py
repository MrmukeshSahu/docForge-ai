import docx

class DocumentParser:
    def __init__(self, doc):
        self.doc = doc

    def count_images(self):
        image_count = 0
        for p in self.doc.paragraphs:
            for run in p.runs:
                drawings = run._element.findall('.//w:drawing', namespaces=run._element.nsmap)
                image_count += len(drawings)
        return image_count

    def extract_tables(self):
        return self.doc.tables

    def extract_elements(self):
        for i, p in enumerate(self.doc.paragraphs):
            text = " ".join(p.text.split())
            if not text:
                continue
            
            # Heuristic to check if paragraph is mostly bold
            bold_chars = sum(len(run.text) for run in p.runs if run.bold)
            total_chars = len(text)
            is_bold = (bold_chars / total_chars > 0.5) if total_chars > 0 else False
            
            # Heuristic to check for monospace font (Code Block detection)
            is_monospace = False
            for run in p.runs:
                if run.font.name and run.font.name.lower() in ["courier", "courier new", "consolas", "monospace", "lucida console"]:
                    is_monospace = True
                    break

            # Heuristic for Blockquote (indented paragraph)
            left_indent = p.paragraph_format.left_indent
            is_indented = (left_indent is not None and left_indent.pt > 18)

            yield {
                'index': i,
                'text': text,
                'is_bold': is_bold,
                'is_monospace': is_monospace,
                'is_indented': is_indented,
                'alignment': p.alignment,
                'paragraph_obj': p,
                'char_count': total_chars,
                'word_count': len(text.split())
            }

