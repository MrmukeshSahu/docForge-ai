import docx
import re
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from presets import BUILTIN_PRESETS, StylePreset

class DocumentFormatter:
    def __init__(self, doc, preset: StylePreset = None):
        self.doc = doc
        self.preset = preset if preset is not None else BUILTIN_PRESETS["classic_book"]
        
        self.page_width = Cm(21.59) # Default Letter width
        self.left_margin = Cm(self.preset.left_margin_cm)
        self.right_margin = Cm(self.preset.right_margin_cm)
        self.heading1_count = 0
        self._setup_page_layout()
        self._format_tables()

    def _setup_page_layout(self):
        for section in self.doc.sections:
            if section.page_width:
                self.page_width = section.page_width
            section.top_margin = Cm(self.preset.top_margin_cm)
            section.bottom_margin = Cm(self.preset.bottom_margin_cm)
            section.left_margin = Cm(self.preset.left_margin_cm)
            section.right_margin = Cm(self.preset.right_margin_cm)
            section.gutter = Cm(0)
            
            # 2-Column Layout Support via OXML
            if getattr(self.preset, 'two_column_layout', False):
                sectPr = section._sectPr
                cols = parse_xml(r'<w:cols %s w:num="2" w:space="720"/>' % nsdecls('w'))
                sectPr.append(cols)

            # Pagination & Header/Footer excluding first page
            section.different_first_page_header_footer = True

            
            # Footer: Page Numbers
            if self.preset.add_page_numbers:
                footer = section.footer
                if not footer.paragraphs:
                    footer.add_paragraph()
                footer.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                self._add_page_number(footer.paragraphs[0])
            
            # Header: Book Title & STYLEREF
            if self.preset.add_chapter_header:
                header = section.header
                if not header.paragraphs:
                    header.add_paragraph()
                hp = header.paragraphs[0]
                hp.text = ""
                
                hp.add_run("Manuscript Title\t\t")
                # STYLEREF OXML for dynamic current chapter
                fldChar1 = parse_xml(r'<w:fldChar %s w:fldCharType="begin"/>' % nsdecls('w'))
                instrText = parse_xml(r'<w:instrText %s xml:space="preserve"> STYLEREF "Heading 1" \* MERGEFORMAT </w:instrText>' % nsdecls('w'))
                fldChar2 = parse_xml(r'<w:fldChar %s w:fldCharType="separate"/>' % nsdecls('w'))
                fldChar3 = parse_xml(r'<w:fldChar %s w:fldCharType="end"/>' % nsdecls('w'))
                
                run_style = hp.add_run()
                r_xml = run_style._r
                r_xml.append(fldChar1)
                r_xml.append(instrText)
                r_xml.append(fldChar2)
                r_xml.append(fldChar3)

    def _add_page_number(self, paragraph):
        p_xml = paragraph._p
        fldChar1 = parse_xml(r'<w:fldChar %s w:fldCharType="begin"/>' % nsdecls('w'))
        instrText = parse_xml(r'<w:instrText %s xml:space="preserve">PAGE</w:instrText>' % nsdecls('w'))
        fldChar2 = parse_xml(r'<w:fldChar %s w:fldCharType="separate"/>' % nsdecls('w'))
        fldChar3 = parse_xml(r'<w:fldChar %s w:fldCharType="end"/>' % nsdecls('w'))
        
        run = paragraph.add_run()
        r_xml = run._r
        r_xml.append(fldChar1)
        r_xml.append(instrText)
        r_xml.append(fldChar2)
        r_xml.append(fldChar3)

    def _format_tables(self):
        try:
            from table_visualizer import TableStylingEngine
            hex_col = "%02x%02x%02x" % self.preset.primary_color_rgb
            TableStylingEngine.apply_zebra_striping(self.doc, primary_color_hex=hex_col)
        except Exception:
            for table in self.doc.tables:
                try:
                    table.style = self.preset.table_style
                except KeyError:
                    table.style = 'Table Grid'


    def _set_font(self, p, size, bold=None, italic=None, font_name=None, color_rgb=None):
        target_font = font_name if font_name else self.preset.font_family
        target_color = color_rgb if color_rgb else self.preset.primary_color_rgb
        
        for run in p.runs:
            font = run.font
            font.name = target_font
            font.size = Pt(size)
            font.color.rgb = RGBColor(*target_color)
            if bold is True: font.bold = True
            if italic is True: font.italic = True
                
        # Auto-Fit Media Logic
        max_width = self.page_width - self.left_margin - self.right_margin
        for shape in self.doc.inline_shapes:
            if shape.width > max_width:
                aspect_ratio = shape.height / shape.width
                shape.width = int(max_width)
                shape.height = int(max_width * aspect_ratio)

    def _add_toc(self, paragraph):
        p = paragraph._p
        toc_xml = parse_xml(r'''
        <w:sdt %s>
            <w:sdtPr>
                <w:docPartObj>
                    <w:docPartGallery w:val="Table of Contents"/>
                    <w:docPartUnique/>
                </w:docPartObj>
            </w:sdtPr>
            <w:sdtContent>
                <w:p>
                    <w:pPr>
                        <w:pStyle w:val="TOCHeading"/>
                    </w:pPr>
                    <w:r><w:t>Table of Contents</w:t></w:r>
                </w:p>
                <w:p>
                    <w:r>
                        <w:fldChar w:fldCharType="begin"/>
                        <w:instrText xml:space="preserve"> TOC \o "1-3" \h \z \u </w:instrText>
                        <w:fldChar w:fldCharType="separate"/>
                    </w:r>
                    <w:r><w:t>Right-click to update field.</w:t></w:r>
                    <w:r><w:fldChar w:fldCharType="end"/></w:r>
                </w:p>
            </w:sdtContent>
        </w:sdt>
        ''' % nsdecls('w'))
        p.addnext(toc_xml)

    def _apply_paragraph_style(self, p, style_type):
        pf = p.paragraph_format
        pf.widow_control = True 
        
        if style_type == "Title":
            pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf.first_line_indent = Cm(0)
            self._set_font(p, self.preset.title_size, bold=True, color_rgb=self.preset.heading_color_rgb)
            pf.keep_with_next = True
            
        elif style_type == "Author":
            pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf.first_line_indent = Cm(0)
            self._set_font(p, self.preset.body_size + 2)
            if self.preset.add_toc:
                self._add_toc(p)
            
        elif style_type in ["Heading 1", "References Heading"]:
            self.heading1_count += 1
            if self.heading1_count > 1:
                p.insert_paragraph_before().add_run().add_break(docx.enum.text.WD_BREAK.PAGE)
            
            pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf.first_line_indent = Cm(0)
            pf.space_before = Pt(12)
            pf.space_after = Pt(6)
            pf.keep_with_next = True
            self._set_font(p, self.preset.heading1_size, bold=True, color_rgb=self.preset.heading_color_rgb)
            
        elif style_type in ["Subheading", "Heading 2"]:
            pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf.first_line_indent = Cm(0)
            pf.space_before = Pt(6)
            pf.space_after = Pt(4)
            pf.keep_with_next = True
            self._set_font(p, self.preset.heading2_size, bold=True, color_rgb=self.preset.heading_color_rgb)
            
        elif style_type == "Heading 3":
            pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf.first_line_indent = Cm(0)
            pf.space_before = Pt(4)
            pf.space_after = Pt(2)
            pf.keep_with_next = True
            self._set_font(p, self.preset.heading3_size, bold=True, italic=True, color_rgb=self.preset.heading_color_rgb)
            
        elif style_type == "Caption":
            pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf.first_line_indent = Cm(0)
            self._set_font(p, self.preset.caption_size, italic=True)
            
        elif style_type == "List":
            pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf.first_line_indent = Cm(0)
            pf.left_indent = Cm(1.27)
            pf.line_spacing = self.preset.line_spacing
            self._set_font(p, self.preset.body_size)
            
        elif style_type == "Code Block":
            pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf.first_line_indent = Cm(0)
            pf.left_indent = Cm(1.0)
            pf.line_spacing = 1.15
            pf.space_before = Pt(4)
            pf.space_after = Pt(4)
            self._set_font(p, self.preset.body_size - 1, font_name="Consolas")
            
        elif style_type == "Blockquote":
            pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf.first_line_indent = Cm(0)
            pf.left_indent = Cm(1.5)
            pf.right_indent = Cm(1.5)
            pf.line_spacing = self.preset.line_spacing
            self._set_font(p, self.preset.body_size, italic=True)
            
        elif style_type == "Reference Item":
            pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf.left_indent = Cm(1.27)
            pf.first_line_indent = Cm(-1.27)
            pf.line_spacing = self.preset.line_spacing
            self._set_font(p, self.preset.body_size)
            
        else:
            # Body Paragraph
            align_enum = WD_ALIGN_PARAGRAPH.JUSTIFY if self.preset.body_alignment == "JUSTIFY" else WD_ALIGN_PARAGRAPH.LEFT
            pf.alignment = align_enum
            pf.first_line_indent = Cm(self.preset.first_line_indent_cm)
            pf.line_spacing = self.preset.line_spacing
            pf.keep_together = True
            self._set_font(p, self.preset.body_size)

    def add_review_comment(self, p):
        # Visual DOCX Comment for Low Confidence
        run = p.add_run(" [Review Needed: Please verify formatting]")
        run.font.color.rgb = RGBColor(255, 0, 0)
        run.font.bold = True
        # Add yellow highlight using OXML
        rPr = run._r.get_or_add_rPr()
        shd = parse_xml(r'<w:shd {} w:val="clear" w:color="auto" w:fill="FFFF00"/>'.format(nsdecls('w')))
        rPr.append(shd)

    def format_elements(self, elements_with_classes):
        in_references = False
        fig_counter = 0
        table_counter = 0

        for item in elements_with_classes:
            p = item['paragraph_obj']
            cls = item['classification']
            review_needed = item.get('review_needed', False)
            
            # Caption Auto-Renumbering Feature
            if cls == "Caption":
                raw_text = p.text.strip()
                # Check figure vs table caption
                if re.match(r'^(figure|fig\.)', raw_text, re.IGNORECASE):
                    fig_counter += 1
                    cleaned_body = re.sub(r'^(figure|fig\.)\s*[\d\.]*[:\s]*', '', raw_text, flags=re.IGNORECASE).strip()
                    new_text = f"Figure {fig_counter}: {cleaned_body}" if cleaned_body else f"Figure {fig_counter}"
                    if p.runs:
                        p.runs[0].text = new_text
                        for r in p.runs[1:]:
                            r.text = ""
                elif re.match(r'^(table|tbl\.)', raw_text, re.IGNORECASE):
                    table_counter += 1
                    cleaned_body = re.sub(r'^(table|tbl\.)\s*[\d\.]*[:\s]*', '', raw_text, flags=re.IGNORECASE).strip()
                    new_text = f"Table {table_counter}: {cleaned_body}" if cleaned_body else f"Table {table_counter}"
                    if p.runs:
                        p.runs[0].text = new_text
                        for r in p.runs[1:]:
                            r.text = ""

            if cls == "References Heading":
                in_references = True
            elif in_references and cls == "Body Paragraph":
                cls = "Reference Item"
                
            self._apply_paragraph_style(p, cls)
            
            if review_needed:
                self.add_review_comment(p)

            
    def save(self, output_path):
        self.doc.save(output_path)

