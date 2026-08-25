import docx
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
from docx.shared import Pt, RGBColor

class WatermarkGenerator:
    @staticmethod
    def apply_watermark(doc, watermark_text: str = "CONFIDENTIAL"):
        if not watermark_text or not watermark_text.strip():
            return

        wm_str = watermark_text.strip().upper()

        for idx, section in enumerate(doc.sections, start=1):
            header = section.header
            p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            
            # Official MS Word OpenXML VML Watermark Shape with full shapetype, color2 fill, o:allowincell="f" & 315-degree rotation
            vml_xml = f"""
            <w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" 
                 xmlns:v="urn:schemas-microsoft-com:vml" 
                 xmlns:o="urn:schemas-microsoft-com:office:office">
                <w:rPr>
                    <w:noProof/>
                </w:rPr>
                <w:pict>
                    <v:shapetype id="_x0000_t136" coordsize="21600,21600" o:spt="136" path="m@7,l@8,m@5,l@6,e">
                        <v:formulas>
                            <v:f eqn="sum #0 0 10800"/>
                            <v:f eqn="prod #0 2 1"/>
                            <v:f eqn="sum 21600 0 @1"/>
                            <v:f eqn="sum 0 0 @2"/>
                            <v:f eqn="sum 21600 0 @3"/>
                            <v:f eqn="sum #1 0 10800"/>
                            <v:f eqn="prod #1 2 1"/>
                            <v:f eqn="sum 21600 0 @6"/>
                            <v:f eqn="sum 0 0 @7"/>
                            <v:f eqn="sum 21600 0 @8"/>
                        </v:formulas>
                        <v:path textpathok="t" o:connecttype="custom" o:connectlocs="@9,0;@10,10800;@11,21600;@12,10800" o:connectangles="270,180,90,0"/>
                        <v:textpath on="t" fitshape="t"/>
                        <v:handles>
                            <v:h position="#0,bottomRight" xrange="6000,21600"/>
                        </v:handles>
                        <o:lock v:ext="edit" text="t" version="1.0"/>
                    </v:shapetype>
                    <v:shape id="PowerPlusWaterMarkObject{idx}" type="#_x0000_t136" 
                        style="position:absolute;margin-left:0;margin-top:0;width:420pt;height:200pt;rotation:315;z-index:-1;mso-position-horizontal:center;mso-position-horizontal-relative:margin;mso-position-vertical:center;mso-position-vertical-relative:margin" 
                        o:allowincell="f" fillcolor="#64748b" stroked="f">
                        <v:fill opacity="50%"/>
                        <v:textpath style="font-family:&quot;Calibri&quot;;font-size:1pt;font-weight:bold" string="{wm_str}"/>
                    </v:shape>
                </w:pict>
            </w:r>
            """

            try:
                r_elem = parse_xml(vml_xml)
                p._p.insert(0, r_elem)
            except Exception as e:
                p_wm = header.add_paragraph()
                p_wm.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                run = p_wm.add_run(f"[{wm_str}]")
                run.font.size = Pt(36)
                run.font.bold = True
                run.font.color.rgb = RGBColor(161, 161, 170)
