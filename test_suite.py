import unittest
import os
import shutil
from docx import Document
from mock_generator import create_mock_docx
from document_parser import DocumentParser
from classifier import ParagraphClassifier
from cleaner import DocumentCleaner
from consistency_checker import GrammarConsistencyChecker
from acronym_extractor import AcronymGlossaryGenerator
from presets import PresetManager, StylePreset
from analytics import DocumentAnalytics
from formatter import DocumentFormatter
from exporter import DocumentExporter
from diff_tool import DocumentDiffTool
from history import HistoryManager
from parallel_batch import ParallelBatchEngine

def classify_elements_helper(elements):
    classifier = ParagraphClassifier()
    classified_elements = []
    for el in elements:
        cls_str, conf, review_needed, exp = classifier.classify_with_explanation(el)
        el['classification'] = cls_str
        el['confidence'] = conf
        el['review_needed'] = review_needed
        el['explanation'] = exp
        classified_elements.append(el)
    return classified_elements

class TestDocForgeAI(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        cls.test_dir = os.path.abspath("test_temp_dir")
        os.makedirs(cls.test_dir, exist_ok=True)
        cls.input_docx = os.path.join(cls.test_dir, "test_input.docx")
        create_mock_docx(cls.input_docx, pages=2)

    @classmethod
    def tearDownClass(cls):
        if os.path.exists(cls.test_dir):
            shutil.rmtree(cls.test_dir, ignore_errors=True)

    def test_01_document_parser(self):
        doc = Document(self.input_docx)
        parser = DocumentParser(doc)
        elements = list(parser.extract_elements())
        self.assertIsInstance(elements, list)
        self.assertGreater(len(elements), 0)
        self.assertIn('text', elements[0])

    def test_02_classifier(self):
        doc = Document(self.input_docx)
        elements = list(DocumentParser(doc).extract_elements())
        classified = classify_elements_helper(elements)
        self.assertEqual(len(classified), len(elements))
        for item in classified:
            self.assertIn('classification', item)
            self.assertIn('confidence', item)

    def test_03_cleaner(self):
        doc = Document(self.input_docx)
        doc.add_paragraph("This is a test--with em-dash and  double space.")
        cleaner = DocumentCleaner(doc, bullet_standardization=True, trailing_whitespace=True)
        fixed = cleaner.clean()
        self.assertGreaterEqual(fixed, 0)

    def test_04_consistency_checker(self):
        doc = Document(self.input_docx)
        elements = list(DocumentParser(doc).extract_elements())
        report = GrammarConsistencyChecker.analyze_consistency(elements)
        self.assertIn('dominant_dialect', report)
        self.assertIn('us_spelling_count', report)
        self.assertIn('uk_spelling_count', report)

    def test_05_acronym_extractor(self):
        elements = [
            {'text': 'We use AI and ML along with REST API.'},
            {'text': 'The IEEE standard for DOCX and PDF export.'}
        ]
        acronyms = AcronymGlossaryGenerator.extract_acronyms(elements)
        self.assertIn('AI', acronyms)
        self.assertIn('ML', acronyms)
        self.assertIn('IEEE', acronyms)

    def test_06_presets_manager(self):
        pm = PresetManager()
        presets = pm.list_all_presets()
        self.assertIn('classic_book', presets)
        self.assertIn('academic_ieee', presets)
        preset = pm.get_preset('classic_book')
        self.assertIsInstance(preset, StylePreset)

    def test_07_analytics(self):
        doc = Document(self.input_docx)
        elements = list(DocumentParser(doc).extract_elements())
        classified = classify_elements_helper(elements)
        analytics = DocumentAnalytics.analyze(doc, classified)
        self.assertIn('flesch_score', analytics)
        self.assertIn('total_words', analytics)
        self.assertIn('reading_time_mins', analytics)

    def test_08_formatter(self):
        doc = Document(self.input_docx)
        elements = list(DocumentParser(doc).extract_elements())
        classified = classify_elements_helper(elements)
        pm = PresetManager()
        preset = pm.get_preset('classic_book')
        formatter = DocumentFormatter(doc, preset)
        formatter.format_elements(classified)
        formatted_docx = os.path.join(self.test_dir, "test_output.docx")
        formatter.save(formatted_docx)
        self.assertTrue(os.path.exists(formatted_docx))

    def test_09_exporter(self):
        doc = Document(self.input_docx)
        elements = list(DocumentParser(doc).extract_elements())
        classified = classify_elements_helper(elements)
        exporter = DocumentExporter(classified, title="Test Document")
        
        md_path = os.path.join(self.test_dir, "export_test.md")
        txt_path = os.path.join(self.test_dir, "export_test.txt")
        html_path = os.path.join(self.test_dir, "export_test.html")
        
        exporter.export_markdown(md_path)
        exporter.export_txt(txt_path)
        exporter.export_html(html_path)
        
        self.assertTrue(os.path.exists(md_path))
        self.assertTrue(os.path.exists(txt_path))
        self.assertTrue(os.path.exists(html_path))

    def test_10_diff_tool(self):
        doc = Document(self.input_docx)
        elements = list(DocumentParser(doc).extract_elements())
        classified = classify_elements_helper(elements)
        diff_path = os.path.join(self.test_dir, "diff_test.html")
        DocumentDiffTool.generate_html_diff(elements, classified, diff_path)
        self.assertTrue(os.path.exists(diff_path))

    def test_11_history_manager(self):
        history_file = os.path.join(self.test_dir, "test_history.json")
        hm = HistoryManager(history_file=history_file)
        hm.add_entry("test_file.docx", "test_output.docx", 1.2, 100, "classic_book")
        history = hm.get_history()
        self.assertEqual(len(history), 1)

    def test_12_parallel_batch(self):
        batch_out = os.path.join(self.test_dir, "batch_out")
        res = ParallelBatchEngine.process_batch_parallel(input_dir=self.test_dir, output_dir=batch_out, preset_id='modern_corporate', export_format='docx')
        self.assertGreater(len(res), 0)
        self.assertEqual(res[0]['status'], 'Success')

if __name__ == '__main__':
    unittest.main()
