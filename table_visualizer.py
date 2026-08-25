import docx
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
from docx.shared import Pt, Inches, RGBColor
import os
import tempfile

class TableStylingEngine:
    @staticmethod
    def apply_zebra_striping(doc, primary_color_hex: str = "1E293B"):
        for table in doc.tables:
            # Skip 1x1 layout container tables (e.g. Cover Page frames)
            if len(table.rows) <= 1 and len(table.columns) <= 1:
                continue
            try:
                table.style = 'Table Grid'
            except Exception:
                pass
            for idx, row in enumerate(table.rows):

                if idx == 0:
                    # Header row styling
                    trPr = row._tr.get_or_add_trPr()
                    trPr.append(parse_xml(r'<w:tblHeader %s/>' % nsdecls('w')))
                    for cell in row.cells:
                        shd = parse_xml(r'<w:shd %s w:fill="%s"/>' % (nsdecls('w'), primary_color_hex))
                        cell._tc.get_or_add_tcPr().append(shd)
                        for p in cell.paragraphs:
                            for r in p.runs:
                                r.font.bold = True
                                r.font.color.rgb = RGBColor(255, 255, 255)
                elif idx % 2 == 1:
                    # Alternate Zebra row shading
                    for cell in row.cells:
                        shd = parse_xml(r'<w:shd %s w:fill="F1F5F9"/>' % nsdecls('w'))
                        cell._tc.get_or_add_tcPr().append(shd)

    @staticmethod
    def generate_chart_for_numeric_tables(doc) -> int:
        """
        Scans numeric tables and generates offline Matplotlib charts embedded in the document.
        """
        charts_created = 0
        try:
            import matplotlib
            matplotlib.use('Agg')
            import matplotlib.pyplot as plt

            for table_idx, table in enumerate(doc.tables):
                if len(table.rows) < 3 or len(table.columns) < 2:
                    continue
                
                # Check if columns contain numbers
                headers = [cell.text.strip() for cell in table.rows[0].cells]
                labels = []
                values = []

                for row in table.rows[1:]:
                    cells_text = [cell.text.strip() for cell in row.cells]
                    if len(cells_text) >= 2:
                        labels.append(cells_text[0][:15])
                        try:
                            val = float(cells_text[1].replace(',', '').replace('$', ''))
                            values.append(val)
                        except ValueError:
                            pass

                if len(values) >= 2 and len(labels) == len(values):
                    # Generate Matplotlib chart locally
                    fig, ax = plt.subplots(figsize=(6, 3.2))
                    ax.bar(labels, values, color='#3B82F6', edgecolor='#1D4ED8')
                    ax.set_title(f"Table Data Chart ({headers[1] if len(headers)>1 else 'Metrics'})", fontsize=11, fontweight='bold', color='#0F172A')
                    ax.set_ylabel(headers[1] if len(headers)>1 else 'Value', fontsize=9)
                    plt.xticks(rotation=25, ha='right', fontsize=8)
                    plt.tight_layout()

                    tmp_chart_path = os.path.join(tempfile.gettempdir(), f"table_chart_{table_idx}_{int(os.getpid())}.png")
                    plt.savefig(tmp_chart_path, dpi=150)
                    plt.close(fig)

                    # Embed chart after table
                    p_chart = doc.add_paragraph()
                    p_chart.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                    p_chart.paragraph_format.space_before = Pt(12)
                    p_chart.paragraph_format.space_after = Pt(4)
                    run_img = p_chart.add_run()
                    run_img.add_picture(tmp_chart_path, width=Inches(5.0))

                    p_cap = doc.add_paragraph()
                    p_cap.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                    p_cap.paragraph_format.space_after = Pt(18)
                    r_cap = p_cap.add_run(f"Figure: Visual data summary for {headers[0] if headers else 'Table'}")
                    r_cap.font.size = Pt(9.5)
                    r_cap.font.italic = True
                    r_cap.font.color.rgb = RGBColor(100, 116, 139)

                    charts_created += 1

        except Exception as e:
            print(f"[!] Chart visualizer error: {e}")

        return charts_created
