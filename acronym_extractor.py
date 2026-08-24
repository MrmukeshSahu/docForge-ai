import re
from typing import List, Dict, Any
import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

COMMON_GLOSSARY = {
    "IEEE": "Institute of Electrical and Electronics Engineers",
    "OXML": "Open Office XML Standard",
    "REST": "Representational State Transfer Architecture",
    "NLP": "Natural Language Processing",
    "API": "Application Programming Interface",
    "DOCX": "Microsoft Word XML Document Standard",
    "EPUB": "Electronic Publication eBook Standard",
    "HTML": "HyperText Markup Language",
    "PDF": "Portable Document Format",
    "ML": "Machine Learning",
    "AI": "Artificial Intelligence",
    "TOC": "Table of Contents",
    "WPM": "Words Per Minute",
    "URL": "Uniform Resource Locator"
}

class AcronymGlossaryGenerator:
    @staticmethod
    def extract_acronyms(elements: List[Dict[str, Any]]) -> Dict[str, str]:
        found_acronyms = {}
        for item in elements:
            text = item.get('text', '')
            tokens = re.findall(r'\b[A-Z]{2,6}\b', text)
            for tok in tokens:
                if tok in COMMON_GLOSSARY:
                    found_acronyms[tok] = COMMON_GLOSSARY[tok]
                elif tok not in found_acronyms:
                    found_acronyms[tok] = f"Abbreviation / Technical Term ({tok})"
        return found_acronyms

    @staticmethod
    def insert_acronyms_table(doc, acronyms: Dict[str, str]):
        if not acronyms:
            return

        p_head = doc.add_paragraph()
        p_head.paragraph_format.space_before = Pt(18)
        p_head.paragraph_format.space_after = Pt(6)
        r_head = p_head.add_run("List of Acronyms & Abbreviations")
        r_head.font.size = Pt(14)
        r_head.font.bold = True
        r_head.font.color.rgb = RGBColor(15, 23, 42)

        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Acronym"
        hdr_cells[1].text = "Expanded Definition / Meaning"
        hdr_cells[0].paragraphs[0].runs[0].font.bold = True
        hdr_cells[1].paragraphs[0].runs[0].font.bold = True

        for acr, desc in sorted(acronyms.items()):
            row_cells = table.add_row().cells
            row_cells[0].text = acr
            row_cells[1].text = desc
            row_cells[0].paragraphs[0].runs[0].font.bold = True

        doc.add_paragraph().paragraph_format.space_after = Pt(12)
