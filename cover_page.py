import docx
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import time

class CoverPageGenerator:
    @staticmethod
    def insert_cover_page(doc, title: str = "Document Title", author: str = "docForge Publication Engine", preset_id: str = "classic_book"):
        first_p = doc.paragraphs[0] if doc.paragraphs else doc.add_paragraph()
        
        # Clean title string
        formatted_title = title.replace('_', ' ').replace('-', ' ').strip()
        if not formatted_title or formatted_title.lower() in ["mock input", "input", "document", "chapter 1"]:
            formatted_title = "OFFICIAL PUBLICATION MANUSCRIPT"
        else:
            formatted_title = formatted_title.upper()

        # Preset-specific color themes
        if preset_id == "modern_corporate":
            bg_fill = "0F172A"       # Deep Executive Dark
            accent_hex = "2563EB"    # Vivid Blue
            gold_hex = "60A5FA"      # Light Blue Accent
            card_bg = "1E293B"       # Slate Card
            card_border = "334155"
        elif preset_id == "academic_ieee":
            bg_fill = "0F172A"       # Charcoal Black
            accent_hex = "38BDF8"    # Cyan
            gold_hex = "38BDF8"
            card_bg = "1E293B"
            card_border = "334155"
        else: # classic_book / default
            bg_fill = "0F172A"       # Deep Slate Navy
            accent_hex = "F59E0B"    # Luxury Gold
            gold_hex = "F59E0B"
            card_bg = "1E293B"       # Dark Slate Card
            card_border = "334155"

        # Create Full-Page Master Cover Table before first_p
        tbl = doc.add_table(rows=1, cols=1)
        tbl.autofit = False
        first_p._p.addprevious(tbl._tbl)
        
        cell = tbl.cell(0, 0)
        cell.width = Cm(17.5)
        
        # Apply full cell background shading
        tcPr = cell._element.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg_fill}"/>')
        tcPr.append(shd)
        
        # Cell Margins / Padding (Top 40pt, Bottom 40pt, Left 20pt, Right 20pt)
        tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="800" w:type="dxa"/><w:bottom w:w="800" w:type="dxa"/><w:left w:w="400" w:type="dxa"/><w:right w:w="400" w:type="dxa"/></w:tcMar>')
        tcPr.append(tcMar)

        # 1. Top Edition Badge
        p_top = cell.paragraphs[0]
        p_top.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_top.paragraph_format.space_before = Pt(36)
        p_top.paragraph_format.space_after = Pt(40)
        
        r_top = p_top.add_run("❖   PUBLICATION & RESEARCH STUDIO EDITION   ❖")
        r_top.font.name = 'Arial'
        r_top.font.size = Pt(11)
        r_top.font.bold = True
        r_top.font.color.rgb = RGBColor(245, 158, 11) if preset_id == "classic_book" else RGBColor(96, 165, 250)

        # 2. Main Title Paragraph
        p_title = cell.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_title.paragraph_format.space_before = Pt(20)
        p_title.paragraph_format.space_after = Pt(20)
        
        r_title = p_title.add_run(formatted_title)
        r_title.font.name = 'Georgia' if preset_id == "classic_book" else 'Arial'
        r_title.font.size = Pt(30)
        r_title.font.bold = True
        r_title.font.color.rgb = RGBColor(255, 255, 255) # Pure White Text

        # 3. Accent Divider Line
        p_rule = cell.add_paragraph()
        p_rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_rule.paragraph_format.space_after = Pt(20)
        pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="24" w:space="1" w:color="{accent_hex}"/></w:pBdr>')
        p_rule._p.get_or_add_pPr().append(pBdr)

        # 4. Subtitle Paragraph
        p_sub = cell.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_sub.paragraph_format.space_after = Pt(60)
        
        r_sub = p_sub.add_run("ACADEMIC & PROFESSIONAL MANUSCRIPT SPECIFICATION")
        r_sub.font.name = 'Arial'
        r_sub.font.size = Pt(10.5)
        r_sub.font.bold = True
        r_sub.font.color.rgb = RGBColor(148, 163, 184) # Slate Soft Grey

        # 5. Inner Metadata Card Table
        tbl_meta = cell.add_table(rows=1, cols=1)
        tbl_meta.autofit = False
        
        c_meta = tbl_meta.cell(0, 0)
        c_meta.width = Cm(16.0)
        tcPr_m = c_meta._element.get_or_add_tcPr()
        shd_m = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{card_bg}"/>')
        tcPr_m.append(shd_m)
        
        # Border for inner metadata card
        tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:top w:val="single" w:sz="12" w:space="0" w:color="{card_border}"/><w:left w:val="single" w:sz="36" w:space="0" w:color="{accent_hex}"/><w:bottom w:val="single" w:sz="12" w:space="0" w:color="{card_border}"/><w:right w:val="single" w:sz="12" w:space="0" w:color="{card_border}"/></w:tcBorders>')
        tcPr_m.append(tcBorders)

        p_m = c_meta.paragraphs[0]
        p_m.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_m.paragraph_format.space_before = Pt(16)
        p_m.paragraph_format.space_after = Pt(16)
        p_m.paragraph_format.left_indent = Cm(0.5)

        r_m1 = p_m.add_run("AUTHOR / PUBLISHER: ")
        r_m1.font.name = 'Arial'
        r_m1.font.size = Pt(10.5)
        r_m1.font.bold = True
        r_m1.font.color.rgb = RGBColor(248, 250, 252) # White

        r_m2 = p_m.add_run(f"{author}\n")
        r_m2.font.name = 'Arial'
        r_m2.font.size = Pt(10.5)
        r_m2.font.color.rgb = RGBColor(203, 213, 225) # Soft Blue Grey

        r_m3 = p_m.add_run("DATE OF PUBLICATION: ")
        r_m3.font.name = 'Arial'
        r_m3.font.size = Pt(10.5)
        r_m3.font.bold = True
        r_m3.font.color.rgb = RGBColor(248, 250, 252)

        r_m4 = p_m.add_run(f"{time.strftime('%B %d, %Y')}\n")
        r_m4.font.name = 'Arial'
        r_m4.font.size = Pt(10.5)
        r_m4.font.color.rgb = RGBColor(203, 213, 225)

        r_m5 = p_m.add_run("DOCUMENT VERIFICATION: ")
        r_m5.font.name = 'Arial'
        r_m5.font.size = Pt(10.5)
        r_m5.font.bold = True
        r_m5.font.color.rgb = RGBColor(248, 250, 252)

        r_m6 = p_m.add_run("Verified Studio Master Output (100% Offline ML Formatted)")
        r_m6.font.name = 'Arial'
        r_m6.font.size = Pt(10.5)
        r_m6.font.color.rgb = RGBColor(203, 213, 225)

        # Spacing after metadata card inside master cell
        p_bot_sp = cell.add_paragraph()
        p_bot_sp.paragraph_format.space_before = Pt(40)
        p_bot_sp.paragraph_format.space_after = Pt(20)

        # Page Break after full-page cover page
        first_p.paragraph_format.page_break_before = True
