from docx import Document
from docx.shared import Pt
import random
import string

def generate_lorem(words):
    lorem = "Lorem ipsum dolor sit amet, consectetur adipiscing elit, sed do eiusmod tempor incididunt ut labore et dolore magna aliqua. Ut enim ad minim veniam, quis nostrud exercitation ullamco laboris nisi ut aliquip ex ea commodo consequat. Duis aute irure dolor in reprehenderit in voluptate velit esse cillum dolore eu fugiat nulla pariatur. Excepteur sint occaecat cupidatat non proident, sunt in culpa qui officia deserunt mollit anim id est laborum."
    lorem_words = lorem.split()
    output = []
    while len(output) < words:
        output.extend(lorem_words)
    return " ".join(output[:words])

def create_mock_docx(filename, pages=5):
    doc = Document()
    
    # Title
    p = doc.add_paragraph("THE IMPACT OF MACHINE LEARNING ON PUBLISHING")
    
    # Author
    p = doc.add_paragraph("By John Doe")
    
    # Body
    doc.add_paragraph(generate_lorem(50))
    
    for chapter in range(1, int(pages) + 1):
        # Heading 1
        p = doc.add_paragraph(f"Chapter {chapter}: Introduction to Automation")
        p.runs[0].bold = True # Simulate user making it bold manually
        
        # Body
        doc.add_paragraph(generate_lorem(150))
        
        # Subheading
        doc.add_paragraph(f"{chapter}.1 Historical Context")
        
        # Body
        doc.add_paragraph(generate_lorem(100))
        
        # List
        doc.add_paragraph("* First important point regarding the context.")
        doc.add_paragraph("* Second point which is also very important.")
        
        # Caption
        doc.add_paragraph(f"Figure {chapter}.1: A generic placeholder for an image.")
        
        # Body
        doc.add_paragraph(generate_lorem(200))
        
    doc.add_paragraph("References")
    doc.add_paragraph("1. Doe, J. (2023). Publishing in the AI Era. Tech Press.")
    doc.add_paragraph("2. Smith, A. (2022). Automated Document Analysis. Academic Journal.")
    
    doc.save(filename)
    print(f"Mock document saved as {filename}")

if __name__ == "__main__":
    create_mock_docx("mock_input.docx", pages=10) # 10 chapters to have enough text
