import customtkinter as ctk
from tkinter import filedialog, messagebox, ttk
import threading
import time
import os
import sys

try:
    from tkinterdnd2 import TkinterDnD, DND_FILES
    has_dnd = True
except ImportError:
    has_dnd = False

from cleaner import DocumentCleaner
from document_parser import DocumentParser
from classifier import ParagraphClassifier
from formatter import DocumentFormatter
from presets import PresetManager, StylePreset, BUILTIN_PRESETS
from analytics import DocumentAnalytics
from exporter import DocumentExporter
from history import HistoryManager

# Dynamic Theme Design Tokens (Tuple Format: ("LightModeColor", "DarkModeColor"))
APP_BG = ("#F8FAFC", "#0F172A")          # Main workspace background
SIDEBAR_BG = ("#F1F5F9", "#1E293B")      # Sidebar background
HEADER_BG = ("#F8FAFC", "#0F172A")       # Top header seamless matching workspace
CARD_BG = ("#FFFFFF", "#1E293B")         # Elevate cards: Crisp white in Light, Slate Dark in Dark
CARD_BORDER = ("#E2E8F0", "#334155")     # Card subtle borders
TEXT_MAIN = ("#0F172A", "#F8FAFC")       # Main text: Slate black in Light, White in Dark
TEXT_MUTED = ("#64748B", "#94A3B8")      # Subtext: Cool grey in Light, Slate grey in Dark
ACCENT_BLUE = ("#2563EB", "#3B82F6")     # Primary button accent
ACCENT_GREEN = ("#059669", "#10B981")    # Action button accent
ACCENT_HOVER = ("#1D4ED8", "#2563EB")    # Button hover
INPUT_BG = ("#FFFFFF", "#0F172A")        # Textbox / Input box background

# Set initial appearance mode
ctk.set_appearance_mode("Dark")
ctk.set_default_color_theme("blue")

class TkinterDnDApp(ctk.CTk):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        if has_dnd:
            try:
                self.TkdndVersion = TkinterDnD._require(self)
            except Exception:
                pass

class App(TkinterDnDApp if has_dnd else ctk.CTk):
    def __init__(self):
        super().__init__()

        self.title("docForge Studio — Intelligent Document Layout Engine")
        self.geometry("1150x740")
        self.minsize(800, 520)
        self.resizable(True, True)
        self.configure(fg_color=APP_BG)


        # Application State Variables
        self.input_file = ""
        self.output_file = ""
        self.batch_input_dir = ""
        self.batch_output_dir = ""
        self.preset_manager = PresetManager()
        self.active_preset_id = "classic_book"
        self.history_manager = HistoryManager()
        self.cancel_requested = False
        
        # Parsed Elements & Inspection Data
        self.doc_obj = None
        self.classified_elements = []
        self.tables_count = 0
        self.images_count = 0
        self.analytics_data = {}

        # Advanced Cleaner & Export Settings
        self.opt_title_case = ctk.BooleanVar(value=False)
        self.opt_bullet_clean = ctk.BooleanVar(value=True)
        self.opt_trailing_clean = ctk.BooleanVar(value=True)
        self.opt_cover_page = ctk.BooleanVar(value=False)
        self.opt_zip_export = ctk.BooleanVar(value=False)
        self.opt_generate_diff = ctk.BooleanVar(value=True)
        self.opt_watermark = ctk.StringVar(value="")
        self.opt_acronyms = ctk.BooleanVar(value=False)
        self.opt_parallel_batch = ctk.BooleanVar(value=True)
        self.opt_export_format = ctk.StringVar(value="all")
        self.watcher_obj = None



        # Root window Grid configuration for 100% Fullscreen Stretch
        self.grid_rowconfigure(0, weight=1)       # Main workspace expands 100% vertically
        self.grid_rowconfigure(1, weight=0)       # Status bar fixed height
        self.grid_columnconfigure(0, weight=0)    # Sidebar fixed width (220px)
        self.grid_columnconfigure(1, weight=1)    # Main container expands 100% horizontally!

        # Build Integrated Main Layout
        self._build_sidebar()
        self._build_main_content()
        self._build_status_bar()

        # Initialize Treeview Styles
        self._update_treeview_theme(dark=True)

        # Show default view
        self.show_view("single_format")

    # -------------------------------------------------------------
    # UI BUILDERS
    # -------------------------------------------------------------
    def _build_sidebar(self):
        self.sidebar_frame = ctk.CTkFrame(self, width=220, corner_radius=0, fg_color=SIDEBAR_BG)
        self.sidebar_frame.grid(row=0, column=0, rowspan=2, sticky="nsew")
        self.sidebar_frame.grid_propagate(False)

        # Brand / Logo Header
        self.logo_label = ctk.CTkLabel(
            self.sidebar_frame, 
            text="⚡ docForge Studio", 
            font=ctk.CTkFont(size=20, weight="bold"),
            text_color=ACCENT_BLUE
        )
        self.logo_label.pack(anchor="w", padx=20, pady=(25, 20))

        # Nav Buttons Definition
        self.nav_buttons = {}
        nav_items = [
            ("single_format", "📌 Format Document"),
            ("inspector", "🔍 Paragraph Inspector"),
            ("presets", "🎨 Style Presets"),
            ("batch", "📁 Batch Processing"),
            ("analytics", "📊 Analytics"),
            ("cleaner", "🧹 Cleaner Settings"),
            ("console", "🖥️ Console Logs"),
            ("history", "📜 Processing History"),
            ("settings", "⚙️ System Specs")
        ]

        for nav_id, label_text in nav_items:
            btn = ctk.CTkButton(
                self.sidebar_frame,
                text=label_text,
                anchor="w",
                height=38,
                corner_radius=8,
                fg_color="transparent",
                text_color=TEXT_MAIN,
                hover_color=CARD_BG,
                font=ctk.CTkFont(size=13, weight="normal"),
                command=lambda nid=nav_id: self.show_view(nid)
            )
            btn.pack(fill="x", padx=12, pady=2)
            self.nav_buttons[nav_id] = btn

        # Sidebar Footer Status Info
        self.sb_footer = ctk.CTkFrame(self.sidebar_frame, fg_color="transparent")
        self.sb_footer.pack(side="bottom", fill="x", padx=15, pady=20)
        
        self.model_status_lbl = ctk.CTkLabel(
            self.sb_footer, 
            text="● ML Classifier Ready", 
            font=ctk.CTkFont(size=11),
            text_color=ACCENT_GREEN
        )
        self.model_status_lbl.pack(anchor="w")

    def _build_main_content(self):
        self.main_container = ctk.CTkFrame(self, corner_radius=0, fg_color=APP_BG)
        self.main_container.grid(row=0, column=1, sticky="nsew")

        # Top Header INSIDE main container for seamless integration
        self.top_header = ctk.CTkFrame(self.main_container, height=55, corner_radius=0, fg_color=APP_BG)
        self.top_header.pack(side="top", fill="x")

        self.header_title = ctk.CTkLabel(
            self.top_header, 
            text="Format Document", 
            font=ctk.CTkFont(size=18, weight="bold"),
            text_color=TEXT_MAIN
        )
        self.header_title.pack(side="left", padx=25, pady=15)

        self.theme_switch = ctk.CTkSwitch(
            self.top_header,
            text="Dark Mode",
            command=self.toggle_theme,
            progress_color=ACCENT_BLUE[1],
            text_color=TEXT_MAIN
        )
        self.theme_switch.pack(side="right", padx=25, pady=15)
        self.theme_switch.select()

        # Divider line
        self.header_divider = ctk.CTkFrame(self.main_container, height=1, fg_color=CARD_BORDER)
        self.header_divider.pack(side="top", fill="x")


        # Dictionary holding all view frames
        self.views = {}

        self.views["single_format"] = self._create_single_format_view()
        self.views["inspector"] = self._create_inspector_view()
        self.views["presets"] = self._create_presets_view()
        self.views["batch"] = self._create_batch_view()
        self.views["analytics"] = self._create_analytics_view()
        self.views["cleaner"] = self._create_cleaner_view()
        self.views["console"] = self._create_console_view()
        self.views["history"] = self._create_history_view()
        self.views["settings"] = self._create_settings_view()

    def _build_status_bar(self):
        self.status_bar = ctk.CTkFrame(self, height=28, corner_radius=0, fg_color=SIDEBAR_BG)
        self.status_bar.grid(row=1, column=1, sticky="ew")

        self.status_text = ctk.CTkLabel(
            self.status_bar,
            text="Ready | Active Preset: Classic Publication Book | Engine: ML RandomForest",
            font=ctk.CTkFont(size=11),
            text_color=TEXT_MUTED
        )
        self.status_text.pack(side="left", padx=15)


    # -------------------------------------------------------------
    # LOGGING & THEME HELPERS
    # -------------------------------------------------------------
    def toggle_theme(self):
        if self.theme_switch.get() == 1:
            ctk.set_appearance_mode("Dark")
            self._update_treeview_theme(dark=True)
            self.status_text.configure(text=f"Mode: Dark | Active Preset: {self.preset_menu_var.get()} | Engine: ML RandomForest")
        else:
            ctk.set_appearance_mode("Light")
            self._update_treeview_theme(dark=False)
            self.status_text.configure(text=f"Mode: Light | Active Preset: {self.preset_menu_var.get()} | Engine: ML RandomForest")

    def _update_treeview_theme(self, dark: bool):
        style = ttk.Style()
        style.theme_use("clam")
        if dark:
            bg = "#1E293B"
            fg = "#F8FAFC"
            field_bg = "#1E293B"
            heading_bg = "#111827"
            heading_fg = "#94A3B8"
        else:
            bg = "#FFFFFF"
            fg = "#0F172A"
            field_bg = "#FFFFFF"
            heading_bg = "#F1F5F9"
            heading_fg = "#475569"

        style.configure("Treeview", background=bg, foreground=fg, fieldbackground=field_bg, borderwidth=0, rowheight=26)
        style.configure("Treeview.Heading", background=heading_bg, foreground=heading_fg, relief="flat", font=('Segoe UI', 10, 'bold'))
        style.map("Treeview", background=[('selected', '#2563EB')], foreground=[('selected', '#FFFFFF')])

    def log_message(self, message: str):
        timestamp = time.strftime("%H:%M:%S")
        log_line = f"[{timestamp}] {message}\n"
        print(log_line.strip())
        self.after(0, self._append_to_console_textbox, log_line)

    def _append_to_console_textbox(self, log_line: str):
        if hasattr(self, 'console_textbox'):
            self.console_textbox.insert("end", log_line)
            self.console_textbox.see("end")

    def cancel_processing(self):
        self.cancel_requested = True
        self.log_message("🛑 Cancel requested by user. Terminating job...")
        self.single_status_lbl.configure(text="Job Cancelled by user.", text_color="#EF4444")
        self.start_btn.configure(state="normal", text="🚀 START FORMATTING")

    # -------------------------------------------------------------
    # VIEW SWITCHER
    # -------------------------------------------------------------
    def show_view(self, view_id: str):
        for v_id, btn in self.nav_buttons.items():
            if v_id == view_id:
                btn.configure(fg_color=ACCENT_BLUE, text_color="#FFFFFF", font=ctk.CTkFont(weight="bold"))
            else:
                btn.configure(fg_color="transparent", text_color=TEXT_MAIN, font=ctk.CTkFont(weight="normal"))

        for v_id, frame in self.views.items():
            if v_id == view_id:
                frame.pack(fill="both", expand=True, padx=25, pady=20)
            else:
                frame.pack_forget()

        titles = {
            "single_format": "Format Document Studio",
            "inspector": "Interactive Paragraph Inspector",
            "presets": "Style Presets",
            "batch": "Batch Processing Studio",
            "analytics": "Document Readability Analytics",
            "cleaner": "Advanced Cleaner Settings",
            "console": "Real-time Console Logs",
            "history": "Processing Activity History",
            "settings": "System Specifications & Specs"
        }
        self.header_title.configure(text=titles.get(view_id, "docForge Studio"))

    # -------------------------------------------------------------
    # VIEW CREATORS
    # -------------------------------------------------------------
    def _create_single_format_view(self):
        frame = ctk.CTkFrame(self.main_container, fg_color="transparent")

        # Grid container spanning full width & height
        grid_wrapper = ctk.CTkFrame(frame, fg_color="transparent")
        grid_wrapper.pack(fill="both", expand=True, padx=5, pady=5)

        grid_wrapper.grid_columnconfigure(0, weight=65) # Left Column: Large Controls & Dropzone (65% width)
        grid_wrapper.grid_columnconfigure(1, weight=35) # Right Column: Active Preset Specs & Engine Info (35% width)

        grid_wrapper.grid_rowconfigure(0, weight=1)

        # -------------------------------------------------------------
        # LEFT COLUMN (Large Control Cards & Main Actions)
        # -------------------------------------------------------------
        left_box = ctk.CTkFrame(grid_wrapper, fg_color="transparent")
        left_box.grid(row=0, column=0, sticky="nsew", padx=(0, 15))

        # Top Card: Large Source File & Drag & Drop Zone
        src_card = ctk.CTkFrame(left_box, fg_color=CARD_BG, border_color=CARD_BORDER, border_width=1, corner_radius=14)
        src_card.pack(fill="x", pady=(0, 15), ipady=12)

        self.dnd_zone = ctk.CTkFrame(src_card, fg_color=SIDEBAR_BG, corner_radius=10, height=120, cursor="hand2")
        self.dnd_zone.pack(fill="x", padx=20, pady=16)
        self.dnd_zone.pack_propagate(False)

        self.dnd_lbl = ctk.CTkLabel(
            self.dnd_zone, 
            text="📥 Click to Browse or Drag & Drop .docx File Here", 
            font=ctk.CTkFont(size=16, weight="bold"),
            text_color=TEXT_MAIN,
            cursor="hand2"
        )
        self.dnd_lbl.pack(expand=True)

        self.dnd_zone.bind("<Button-1>", lambda e: self.select_input())
        self.dnd_lbl.bind("<Button-1>", lambda e: self.select_input())

        if has_dnd:
            try:
                self.dnd_zone.drop_target_register(DND_FILES)
                self.dnd_zone.dnd_bind('<<Drop>>', self._on_drop_single)
            except Exception:
                pass

        # Destination & Options Large Card
        dest_card = ctk.CTkFrame(left_box, fg_color=CARD_BG, border_color=CARD_BORDER, border_width=1, corner_radius=14)
        dest_card.pack(fill="x", pady=10, ipady=15)

        # Row 1: Save Folder & Custom Name Buttons
        d_row = ctk.CTkFrame(dest_card, fg_color="transparent")
        d_row.pack(anchor="center", pady=(15, 6))
        
        ctk.CTkButton(d_row, text="📁 Select Save Folder", command=self.select_output_folder, fg_color=ACCENT_BLUE, hover_color=ACCENT_HOVER, height=38, width=180).pack(side="left", padx=8)
        ctk.CTkButton(d_row, text="✏️ Custom Name", command=self.select_output_file, fg_color=SIDEBAR_BG, hover_color=CARD_BG, text_color=TEXT_MAIN, height=38, width=140).pack(side="left", padx=8)
        
        self.output_lbl = ctk.CTkLabel(dest_card, text="Auto-saves to source folder if unselected", text_color=TEXT_MUTED, font=ctk.CTkFont(size=12))
        self.output_lbl.pack(anchor="center", pady=(0, 12))

        # Row 2: Style Preset Selector
        opt_preset = ctk.CTkFrame(dest_card, fg_color="transparent")
        opt_preset.pack(fill="x", padx=24, pady=6)

        ctk.CTkLabel(opt_preset, text="Style Preset:", text_color=TEXT_MAIN, font=ctk.CTkFont(size=13, weight="bold"), width=110, anchor="w").pack(side="left")
        self.preset_menu_var = ctk.StringVar(value="Classic Publication Book")
        preset_names = [p.name for p in BUILTIN_PRESETS.values()]
        self.preset_dropdown = ctk.CTkOptionMenu(opt_preset, values=preset_names, variable=self.preset_menu_var, command=self._on_preset_dropdown_change, height=36, text_color="#FFFFFF")
        self.preset_dropdown.pack(side="left", fill="x", expand=True)

        # Row 3: Export Format Selector
        opt_export = ctk.CTkFrame(dest_card, fg_color="transparent")
        opt_export.pack(fill="x", padx=24, pady=6)

        ctk.CTkLabel(opt_export, text="Export Format:", text_color=TEXT_MAIN, font=ctk.CTkFont(size=13, weight="bold"), width=110, anchor="w").pack(side="left")
        ctk.CTkOptionMenu(
            opt_export, 
            values=["DOCX Only", "DOCX + EPUB eBook", "DOCX + HTML Audit Report", "All Formats (DOCX, PDF, EPUB, HTML Audit, MD, TXT)"], 
            variable=self.opt_export_format,
            height=36,
            text_color="#FFFFFF"
        ).pack(side="left", fill="x", expand=True)

        # Row 4: Watermark Entry
        opt_wm = ctk.CTkFrame(dest_card, fg_color="transparent")
        opt_wm.pack(fill="x", padx=24, pady=6)
        ctk.CTkLabel(opt_wm, text="💧 Watermark:", text_color=TEXT_MAIN, font=ctk.CTkFont(size=13, weight="bold"), width=110, anchor="w").pack(side="left")
        ctk.CTkEntry(opt_wm, textvariable=self.opt_watermark, placeholder_text="e.g. CONFIDENTIAL / DRAFT / INTERNAL ONLY", height=34).pack(side="left", fill="x", expand=True)

        # Row 5: Phase 2 & 3 Checkboxes
        opt_row2 = ctk.CTkFrame(dest_card, fg_color="transparent")
        opt_row2.pack(anchor="center", pady=(14, 12))

        ctk.CTkCheckBox(opt_row2, text="Auto Cover Page", variable=self.opt_cover_page, text_color=TEXT_MAIN, font=ctk.CTkFont(size=12)).pack(side="left", padx=8)
        ctk.CTkCheckBox(opt_row2, text="1-Click Zip", variable=self.opt_zip_export, text_color=TEXT_MAIN, font=ctk.CTkFont(size=12)).pack(side="left", padx=8)
        ctk.CTkCheckBox(opt_row2, text="Side Diff", variable=self.opt_generate_diff, text_color=TEXT_MAIN, font=ctk.CTkFont(size=12)).pack(side="left", padx=8)
        ctk.CTkCheckBox(opt_row2, text="Acronyms Glossary", variable=self.opt_acronyms, text_color=TEXT_MAIN, font=ctk.CTkFont(size=12)).pack(side="left", padx=8)


        # Large Main Action Buttons Bar
        act_row = ctk.CTkFrame(left_box, fg_color="transparent")
        act_row.pack(anchor="center", pady=20)

        self.start_btn = ctk.CTkButton(
            act_row,
            text="🚀 START FORMATTING",
            font=ctk.CTkFont(size=16, weight="bold"),
            height=52,
            width=280,
            corner_radius=12,
            fg_color=ACCENT_GREEN,
            hover_color="#059669",
            text_color="#FFFFFF",
            command=self.start_single_formatting
        )
        self.start_btn.pack(side="left", padx=10)

        self.cancel_btn = ctk.CTkButton(
            act_row,
            text="🛑 Cancel Job",
            font=ctk.CTkFont(size=15, weight="bold"),
            height=52,
            corner_radius=12,
            fg_color="#DC2626",
            hover_color="#B91C1C",
            text_color="#FFFFFF",
            width=140,
            command=self.cancel_processing
        )
        self.cancel_btn.pack(side="left", padx=10)

        # Progress and Status Box (Permanently Packed & Visible)
        self.progress_bar = ctk.CTkProgressBar(left_box, mode="determinate", height=14, progress_color=ACCENT_GREEN[1], corner_radius=7)
        self.progress_bar.pack(fill="x", padx=20, pady=(15, 6))
        self.progress_bar.set(0)

        self.single_status_lbl = ctk.CTkLabel(left_box, text="Ready to format document", font=ctk.CTkFont(size=13, weight="bold"), text_color=TEXT_MUTED)
        self.single_status_lbl.pack(pady=(0, 10))



        # -------------------------------------------------------------
        # RIGHT COLUMN (Dynamic Preset Visualizer & Engine Status)
        # -------------------------------------------------------------
        right_box = ctk.CTkFrame(grid_wrapper, fg_color="transparent")
        right_box.grid(row=0, column=1, sticky="nsew", padx=(10, 0))

        # Card 1: Active Preset Preview & Specs
        preset_card = ctk.CTkFrame(right_box, fg_color=CARD_BG, border_color=CARD_BORDER, border_width=1, corner_radius=12)
        preset_card.pack(fill="x", pady=(0, 12), ipady=10)

        ctk.CTkLabel(preset_card, text="🎨 Active Preset Specs", font=ctk.CTkFont(size=15, weight="bold"), text_color=ACCENT_BLUE).pack(anchor="w", padx=16, pady=(12, 6))
        
        self.preset_info_lbl = ctk.CTkLabel(
            preset_card, 
            text="", 
            justify="left", 
            font=ctk.CTkFont(size=12), 
            text_color=TEXT_MAIN
        )
        self.preset_info_lbl.pack(anchor="w", padx=16, pady=(0, 10))
        self._update_preset_info_card(self.preset_menu_var.get())

        # Card 2: ML Classifier & Security Specs
        engine_card = ctk.CTkFrame(right_box, fg_color=CARD_BG, border_color=CARD_BORDER, border_width=1, corner_radius=12)
        engine_card.pack(fill="both", expand=True, pady=0, ipady=10)

        ctk.CTkLabel(engine_card, text="⚡ Engine & Privacy Specs", font=ctk.CTkFont(size=15, weight="bold"), text_color=ACCENT_GREEN).pack(anchor="w", padx=16, pady=(12, 6))
        
        specs_text = (
            "• 🔒 100% Offline & Private (No Cloud/API key)\n"
            "• 🧠 RandomForest ML Classifier (SpaCy POS + NLTK)\n"
            "• 🏷️ 10 Element Classes (Title, Author, H1, H2, Code)\n"
            "• 🔢 Automatic Figure & Table Caption Renumbering\n"
            "• 📑 OXML Dynamic STYLEREF Header & PAGE Footers\n"
            "• 📦 Multi-Format (.docx, .pdf, .epub, .html, .md, .zip)"
        )
        ctk.CTkLabel(engine_card, text=specs_text, justify="left", font=ctk.CTkFont(size=11), text_color=TEXT_MUTED).pack(anchor="w", padx=16, pady=5)

        return frame



    def _create_inspector_view(self):
        frame = ctk.CTkFrame(self.main_container, fg_color="transparent")

        # Header Info Bar
        info_bar = ctk.CTkFrame(frame, fg_color=CARD_BG, border_color=CARD_BORDER, border_width=1, corner_radius=10)
        info_bar.pack(fill="x", pady=(0, 10), ipady=5)
        
        ctk.CTkLabel(info_bar, text="Paragraph Classifications & Manual Override Table", font=ctk.CTkFont(size=14, weight="bold"), text_color=TEXT_MAIN).pack(side="left", padx=15)
        ctk.CTkButton(info_bar, text="🔄 Reload Elements", command=self.parse_input_for_inspector, fg_color=ACCENT_BLUE, text_color="#FFFFFF", height=28, width=120).pack(side="right", padx=15)

        # Treeview Container
        tree_container = ctk.CTkFrame(frame, fg_color=CARD_BG, border_color=CARD_BORDER, border_width=1, corner_radius=10)
        tree_container.pack(fill="both", expand=True)

        columns = ("index", "class", "confidence", "text", "explanation")
        self.inspector_tree = ttk.Treeview(tree_container, columns=columns, show="headings", height=15)
        
        self.inspector_tree.heading("index", text="#")
        self.inspector_tree.heading("class", text="Classification")
        self.inspector_tree.heading("confidence", text="Confidence")
        self.inspector_tree.heading("text", text="Paragraph Snippet")
        self.inspector_tree.heading("explanation", text="Classifier Logic")

        self.inspector_tree.column("index", width=40, anchor="center")
        self.inspector_tree.column("class", width=140, anchor="w")
        self.inspector_tree.column("confidence", width=90, anchor="center")
        self.inspector_tree.column("text", width=380, anchor="w")
        self.inspector_tree.column("explanation", width=240, anchor="w")

        scrollbar = ttk.Scrollbar(tree_container, orient="vertical", command=self.inspector_tree.yview)
        self.inspector_tree.configure(yscrollcommand=scrollbar.set)
        
        self.inspector_tree.pack(side="left", fill="both", expand=True, padx=5, pady=5)
        scrollbar.pack(side="right", fill="y", pady=5)

        # Manual Override Bar
        override_bar = ctk.CTkFrame(frame, fg_color=CARD_BG, border_color=CARD_BORDER, border_width=1, corner_radius=10)
        override_bar.pack(fill="x", pady=10, ipady=5)

        ctk.CTkLabel(override_bar, text="Override Selected Item Class:", font=ctk.CTkFont(weight="bold"), text_color=TEXT_MAIN).pack(side="left", padx=15)
        self.override_var = ctk.StringVar(value="Heading 1")
        classes_list = ["Title", "Author", "Heading 1", "Heading 2", "Heading 3", "Body Paragraph", "Caption", "List", "Code Block", "Blockquote", "References Heading"]
        ctk.CTkOptionMenu(override_bar, values=classes_list, variable=self.override_var, text_color="#FFFFFF").pack(side="left", padx=10)
        ctk.CTkButton(override_bar, text="Apply Override", command=self.apply_manual_override, fg_color=ACCENT_GREEN, text_color="#FFFFFF", hover_color="#059669", width=120).pack(side="left", padx=10)

        return frame

    def _create_presets_view(self):
        frame = ctk.CTkFrame(self.main_container, fg_color="transparent")

        # Top Action Bar for JSON Preset Import/Export
        top_bar = ctk.CTkFrame(frame, fg_color=CARD_BG, border_color=CARD_BORDER, border_width=1, corner_radius=10)
        top_bar.pack(fill="x", pady=(0, 10), ipady=5)
        
        ctk.CTkLabel(top_bar, text="🎨 Custom Preset Studio (.dfpreset JSON)", font=ctk.CTkFont(size=14, weight="bold"), text_color=TEXT_MAIN).pack(side="left", padx=15)
        
        ctk.CTkButton(top_bar, text="📥 Import .dfpreset", command=self._gui_import_preset, fg_color=ACCENT_BLUE, text_color="#FFFFFF", height=28, width=130).pack(side="right", padx=10)
        ctk.CTkButton(top_bar, text="📤 Export Active Preset", command=self._gui_export_preset, fg_color=SIDEBAR_BG, text_color=TEXT_MAIN, height=28, width=150).pack(side="right", padx=5)

        # Preset Cards Scrollable
        scroll = ctk.CTkScrollableFrame(frame, fg_color="transparent")
        scroll.pack(fill="both", expand=True)

        for pid, preset in BUILTIN_PRESETS.items():
            card = ctk.CTkFrame(scroll, fg_color=CARD_BG, border_color=CARD_BORDER, border_width=1, corner_radius=12)
            card.pack(fill="x", pady=8, ipady=8)

            c_head = ctk.CTkFrame(card, fg_color="transparent")
            c_head.pack(fill="x", padx=15, pady=5)
            
            ctk.CTkLabel(c_head, text=preset.name, font=ctk.CTkFont(size=15, weight="bold"), text_color=ACCENT_BLUE).pack(side="left")
            
            select_btn = ctk.CTkButton(
                c_head, 
                text="Activate Preset", 
                command=lambda p_id=pid, p_name=preset.name: self._set_active_preset(p_id, p_name),
                fg_color=ACCENT_GREEN, 
                text_color="#FFFFFF",
                height=28, 
                width=120
            )
            select_btn.pack(side="right")

            ctk.CTkLabel(card, text=preset.description, text_color=TEXT_MUTED, font=ctk.CTkFont(size=12)).pack(anchor="w", padx=15, pady=(0, 8))

            details_text = f"Font: {preset.font_family} | Body: {preset.body_size}pt | Title: {preset.title_size}pt | Spacing: {preset.line_spacing}x | Margins: {preset.top_margin_cm}cm | TOC: {'Yes' if preset.add_toc else 'No'}"
            ctk.CTkLabel(card, text=details_text, text_color=TEXT_MAIN, font=ctk.CTkFont(size=11)).pack(anchor="w", padx=15)

        return frame

    def _gui_import_preset(self):
        f_path = filedialog.askopenfilename(title="Import Custom .dfpreset JSON", filetypes=[("JSON Files", "*.json;*.dfpreset"), ("All Files", "*.*")])
        if f_path:
            try:
                p = self.preset_manager.import_preset_json(f_path)
                messagebox.showinfo("Success", f"Successfully imported preset '{p.name}'!")
            except Exception as e:
                messagebox.showerror("Error", f"Failed to import preset: {e}")

    def _gui_export_preset(self):
        f_path = filedialog.asksaveasfilename(title="Export Active Preset as .dfpreset JSON", defaultextension=".json", filetypes=[("JSON Files", "*.json"), ("All Files", "*.*")])
        if f_path:
            try:
                self.preset_manager.export_preset_json(self.active_preset_id, f_path)
                messagebox.showinfo("Success", f"Successfully exported preset to '{f_path}'!")
            except Exception as e:
                messagebox.showerror("Error", f"Failed to export preset: {e}")


    def _create_batch_view(self):
        frame = ctk.CTkFrame(self.main_container, fg_color="transparent")

        card = ctk.CTkFrame(frame, fg_color=CARD_BG, border_color=CARD_BORDER, border_width=1, corner_radius=12)
        card.pack(fill="x", pady=10, ipady=10)

        ctk.CTkLabel(card, text="Batch Directory Formatting", font=ctk.CTkFont(size=15, weight="bold"), text_color=TEXT_MAIN).pack(anchor="w", padx=20, pady=10)

        # Input Dir
        row1 = ctk.CTkFrame(card, fg_color="transparent")
        row1.pack(fill="x", padx=20, pady=5)
        ctk.CTkButton(row1, text="📁 Source Folder", command=self.select_batch_input, fg_color=ACCENT_BLUE, text_color="#FFFFFF", width=150).pack(side="left", padx=(0, 15))
        self.batch_in_lbl = ctk.CTkLabel(row1, text="No folder selected", text_color=TEXT_MUTED)
        self.batch_in_lbl.pack(side="left", fill="x", expand=True)

        # Output Dir
        row2 = ctk.CTkFrame(card, fg_color="transparent")
        row2.pack(fill="x", padx=20, pady=5)
        ctk.CTkButton(row2, text="💾 Output Folder", command=self.select_batch_output, fg_color=ACCENT_BLUE, text_color="#FFFFFF", width=150).pack(side="left", padx=(0, 15))
        self.batch_out_lbl = ctk.CTkLabel(row2, text="No folder selected", text_color=TEXT_MUTED)
        self.batch_out_lbl.pack(side="left", fill="x", expand=True)

        ctk.CTkButton(frame, text="⚡ START BATCH PROCESSING", font=ctk.CTkFont(size=15, weight="bold"), height=45, fg_color=ACCENT_GREEN, text_color="#FFFFFF", command=self.start_batch_formatting).pack(fill="x", pady=20)

        self.batch_status_lbl = ctk.CTkLabel(frame, text="", font=ctk.CTkFont(size=13), text_color=TEXT_MAIN)
        self.batch_status_lbl.pack(pady=5)

        return frame

    def _create_analytics_view(self):
        frame = ctk.CTkFrame(self.main_container, fg_color="transparent")

        self.stat_grid = ctk.CTkFrame(frame, fg_color="transparent")
        self.stat_grid.pack(fill="x", pady=10)

        # Stat cards
        self.card_words = self._create_stat_card(self.stat_grid, "Total Words", "--")
        self.card_readability = self._create_stat_card(self.stat_grid, "Readability Score", "--")
        self.card_grade = self._create_stat_card(self.stat_grid, "Grade Level", "--")
        self.card_time = self._create_stat_card(self.stat_grid, "Reading Time", "--")

        self.card_words.grid(row=0, column=0, padx=5, pady=5, sticky="ew")
        self.card_readability.grid(row=0, column=1, padx=5, pady=5, sticky="ew")
        self.card_grade.grid(row=0, column=2, padx=5, pady=5, sticky="ew")
        self.card_time.grid(row=0, column=3, padx=5, pady=5, sticky="ew")

        self.stat_grid.grid_columnconfigure((0, 1, 2, 3), weight=1)

        # Hierarchy breakdown card
        tree_card = ctk.CTkFrame(frame, fg_color=CARD_BG, border_color=CARD_BORDER, border_width=1, corner_radius=12)
        tree_card.pack(fill="both", expand=True, pady=10)

        ctk.CTkLabel(tree_card, text="Document Hierarchy & Readability Breakdown", font=ctk.CTkFont(size=14, weight="bold"), text_color=TEXT_MAIN).pack(anchor="w", padx=20, pady=10)
        self.analytics_text = ctk.CTkTextbox(tree_card, fg_color=INPUT_BG, text_color=TEXT_MAIN, font=ctk.CTkFont(family="Consolas", size=12))
        self.analytics_text.pack(fill="both", expand=True, padx=15, pady=15)
        self.analytics_text.insert("1.0", "Format or inspect a document to view analytics metrics here.")

        return frame

    def _create_stat_card(self, parent, title, initial_val):
        card = ctk.CTkFrame(parent, fg_color=CARD_BG, border_color=CARD_BORDER, border_width=1, corner_radius=10)
        ctk.CTkLabel(card, text=title, font=ctk.CTkFont(size=11), text_color=TEXT_MUTED).pack(anchor="w", padx=12, pady=(10, 2))
        lbl_val = ctk.CTkLabel(card, text=initial_val, font=ctk.CTkFont(size=18, weight="bold"), text_color=ACCENT_BLUE)
        lbl_val.pack(anchor="w", padx=12, pady=(0, 10))
        card.lbl_val = lbl_val
        return card

    def _create_cleaner_view(self):
        frame = ctk.CTkFrame(self.main_container, fg_color="transparent")

        card = ctk.CTkFrame(frame, fg_color=CARD_BG, border_color=CARD_BORDER, border_width=1, corner_radius=12)
        card.pack(fill="x", pady=10, ipady=15)

        ctk.CTkLabel(card, text="Advanced Typography Cleaning Rules", font=ctk.CTkFont(size=15, weight="bold"), text_color=TEXT_MAIN).pack(anchor="w", padx=20, pady=10)

        ctk.CTkCheckBox(card, text="Standardize Typography (Double hyphens to Em-Dashes, Smart Quotes)", text_color=TEXT_MAIN, state="disabled", variable=ctk.BooleanVar(value=True)).pack(anchor="w", padx=25, pady=8)
        ctk.CTkCheckBox(card, text="Remove Redundant Spaces, Tabs, and Blank Paragraphs", text_color=TEXT_MAIN, state="disabled", variable=ctk.BooleanVar(value=True)).pack(anchor="w", padx=25, pady=8)
        ctk.CTkCheckBox(card, text="Auto-convert Headings to Title Case", text_color=TEXT_MAIN, variable=self.opt_title_case).pack(anchor="w", padx=25, pady=8)
        ctk.CTkCheckBox(card, text="Standardize List Bullet Symbols (*, -, + to •)", text_color=TEXT_MAIN, variable=self.opt_bullet_clean).pack(anchor="w", padx=25, pady=8)
        ctk.CTkCheckBox(card, text="Remove Trailing Whitespaces Across Runs", text_color=TEXT_MAIN, variable=self.opt_trailing_clean).pack(anchor="w", padx=25, pady=8)

        return frame

    def _create_console_view(self):
        frame = ctk.CTkFrame(self.main_container, fg_color="transparent")

        console_card = ctk.CTkFrame(frame, fg_color=CARD_BG, border_color=CARD_BORDER, border_width=1, corner_radius=12)
        console_card.pack(fill="both", expand=True, pady=10)

        c_head = ctk.CTkFrame(console_card, fg_color="transparent")
        c_head.pack(fill="x", padx=15, pady=10)
        
        ctk.CTkLabel(c_head, text="Real-time Execution & Processing Logs", font=ctk.CTkFont(size=14, weight="bold"), text_color=TEXT_MAIN).pack(side="left")
        
        ctk.CTkButton(c_head, text="📋 Copy Logs to Clipboard", command=self.copy_console_logs, fg_color=ACCENT_BLUE, text_color="#FFFFFF", height=28, width=170).pack(side="right")

        self.console_textbox = ctk.CTkTextbox(console_card, fg_color=INPUT_BG, text_color=TEXT_MAIN, font=ctk.CTkFont(family="Consolas", size=11))
        self.console_textbox.pack(fill="both", expand=True, padx=15, pady=15)
        self.console_textbox.insert("1.0", "--- Real-time Execution Console initialized ---\n")

        return frame

    def copy_console_logs(self):
        logs = self.console_textbox.get("1.0", "end")
        self.clipboard_clear()
        self.clipboard_append(logs)
        messagebox.showinfo("Copied", "Console logs copied to clipboard!")

    def _create_history_view(self):
        frame = ctk.CTkFrame(self.main_container, fg_color="transparent")

        history_card = ctk.CTkFrame(frame, fg_color=CARD_BG, border_color=CARD_BORDER, border_width=1, corner_radius=12)
        history_card.pack(fill="both", expand=True, pady=10)

        h_head = ctk.CTkFrame(history_card, fg_color="transparent")
        h_head.pack(fill="x", padx=15, pady=10)
        ctk.CTkLabel(h_head, text="Processing History Log", font=ctk.CTkFont(size=14, weight="bold"), text_color=TEXT_MAIN).pack(side="left")
        ctk.CTkButton(h_head, text="🔄 Refresh Log", command=self.refresh_history, fg_color=ACCENT_BLUE, text_color="#FFFFFF", height=28, width=120).pack(side="right")

        self.history_textbox = ctk.CTkTextbox(history_card, fg_color=INPUT_BG, text_color=TEXT_MAIN, font=ctk.CTkFont(family="Consolas", size=11))
        self.history_textbox.pack(fill="both", expand=True, padx=15, pady=15)
        self.refresh_history()

        return frame

    def refresh_history(self):
        entries = self.history_manager.get_history()
        if hasattr(self, 'history_textbox'):
            self.history_textbox.delete("1.0", "end")
            if not entries:
                self.history_textbox.insert("1.0", "No recent processing history.")
                return

            h_str = ""
            for entry in entries:
                h_str += f"[{entry['timestamp']}] Preset: {entry['preset_used']:<24} | Time: {entry['elapsed_sec']}s | Elements: {entry['elements_count']}\n"
                h_str += f"  Input:  {entry['input_file']}\n"
                h_str += f"  Output: {entry['output_file']}\n"
                h_str += "-" * 85 + "\n"

            self.history_textbox.insert("1.0", h_str)

    def _create_settings_view(self):
        frame = ctk.CTkFrame(self.main_container, fg_color="transparent")

        scroll = ctk.CTkScrollableFrame(frame, fg_color="transparent")
        scroll.pack(fill="both", expand=True)

        # Overview & Best Use Cases Card
        card1 = ctk.CTkFrame(scroll, fg_color=CARD_BG, border_color=CARD_BORDER, border_width=1, corner_radius=12)
        card1.pack(fill="x", pady=8, ipady=10)

        ctk.CTkLabel(card1, text="🎯 Best Use Cases & Document Types", font=ctk.CTkFont(size=15, weight="bold"), text_color=ACCENT_BLUE).pack(anchor="w", padx=20, pady=8)

        use_cases = (
            "• 📚 Books & Novels: Georgia font, 1.25 line spacing, classic book margins, chapter flow & EPUB eBook export.\n"
            "• 🎓 Research Papers (IEEE / APA 7th): Standard 1-inch margins, APA 2.0 double spacing / IEEE compact formatting, figure & table renumbering (Figure 1: ..., Table 1: ...) & bibliography support.\n"
            "• 🏢 Corporate & Business Reports: Executive styling (Calibri/Arial, Navy Blue accent headers, 1.15 line spacing, clean table borders).\n"
            "• 📄 Structured Manuscripts: Unformatted text documents containing Headings (H1, H2), Body paragraphs, Bullet Lists, and Figures/Tables."
        )
        ctk.CTkLabel(card1, text=use_cases, justify="left", text_color=TEXT_MAIN, font=ctk.CTkFont(size=12)).pack(anchor="w", padx=20, pady=5)

        # Core Features Card
        card2 = ctk.CTkFrame(scroll, fg_color=CARD_BG, border_color=CARD_BORDER, border_width=1, corner_radius=12)
        card2.pack(fill="x", pady=8, ipady=10)

        ctk.CTkLabel(card2, text="⚡ Core System Capabilities", font=ctk.CTkFont(size=15, weight="bold"), text_color=ACCENT_GREEN).pack(anchor="w", padx=20, pady=8)

        features_text = (
            "• 🧠 100% Offline ML Classifier: RandomForest + spaCy POS + NLTK NER + Structural Rules.\n"
            "• 📑 10 Element Types Detected: Title, Author, Heading 1, Subheadings, Body, Tables, Figures, Captions, References, Lists.\n"
            "• 🧹 Typography Cleaning: Double spaces, extra blank lines, em-dashes, smart quotes, trailing spaces.\n"
            "• 🔢 Automatic Caption Renumbering: Figure 1: ..., Figure 2: ..., Table 1: ..., Table 2: ...\n"
            "• 📐 Layout & DOM: Table grid borders, image auto-scaling, dynamic header STYLEREF chapter titles, PAGE footers, TOC.\n"
            "• 📤 Multi-Format Export: .DOCX, .PDF, .EPUB eBook, .HTML Audit Report, Markdown, Text.\n"
            "• 🔒 Zero Content Modification: 100% data safety, no internet/cloud API key required, zero text rewriting."
        )
        ctk.CTkLabel(card2, text=features_text, justify="left", text_color=TEXT_MAIN, font=ctk.CTkFont(size=12)).pack(anchor="w", padx=20, pady=5)

        # Known Limitations Card
        card3 = ctk.CTkFrame(scroll, fg_color=CARD_BG, border_color=CARD_BORDER, border_width=1, corner_radius=12)
        card3.pack(fill="x", pady=8, ipady=10)

        ctk.CTkLabel(card3, text="⚠️ System Limitations", font=ctk.CTkFont(size=15, weight="bold"), text_color="#EF4444").pack(anchor="w", padx=20, pady=8)

        limitations_text = (
            "1. 📰 Magazine / Brochure / Flyer Layouts: Floating text boxes and graphic flyer layouts are not supported. Designed for text-first documents.\n"
            "2. 🖼️ Scanned PDF / Image Documents: Requires editable .docx files (No direct scanned image OCR scanner).\n"
            "3. 📊 Complex Merged-Cell Tables: Applies clean borders to standard tables; does not restructure complex multi-row merged XML forms.\n"
            "4. 🌐 Non-English ML Classification: ML model is trained on English NLP patterns. Margins/fonts apply to any language, but non-English headings may require manual inspector override.\n"
            "5. 🌀 Dumped Unstructured Notes: Single continuous text blocks without capitalization require heading line breaks for optimal ML classification."
        )
        ctk.CTkLabel(card3, text=limitations_text, justify="left", text_color=TEXT_MUTED, font=ctk.CTkFont(size=12)).pack(anchor="w", padx=20, pady=5)

        return frame

    def _update_preset_info_card(self, selected_name: str):
        for pid, p in BUILTIN_PRESETS.items():
            if p.name == selected_name:
                txt = f"• Name: {p.name}\n"
                txt += f"• Font Family: {p.font_family}\n"
                txt += f"• Title: {p.title_size}pt | Body: {p.body_size}pt | Captions: {p.caption_size}pt\n"
                txt += f"• Headings: H1 {p.heading1_size}pt, H2 {p.heading2_size}pt, H3 {p.heading3_size}pt\n"
                txt += f"• Line Spacing: {p.line_spacing}x | Indent: {p.first_line_indent_cm}cm\n"
                txt += f"• Margins: Top {p.top_margin_cm}cm, Side {p.left_margin_cm}cm\n"
                txt += f"• 2-Column Mode: {'Yes' if getattr(p, 'two_column_layout', False) else 'No'}\n"
                txt += f"• Chapter Headers: {'Yes' if p.add_chapter_header else 'No'}"
                if hasattr(self, 'preset_info_lbl'):
                    self.preset_info_lbl.configure(text=txt)
                break

    def _set_active_preset(self, pid, pname):
        self.active_preset_id = pid
        self.preset_menu_var.set(pname)
        self._update_preset_info_card(pname)
        self.status_text.configure(text=f"Active Preset: {pname} | Engine: ML RandomForest")
        messagebox.showinfo("Preset Activated", f"Activated Preset: '{pname}'")

    def _on_preset_dropdown_change(self, selected_name):
        for pid, p in BUILTIN_PRESETS.items():
            if p.name == selected_name:
                self.active_preset_id = pid
                self._update_preset_info_card(selected_name)
                self.status_text.configure(text=f"Active Preset: {selected_name} | Engine: ML RandomForest")
                break


    def select_input(self):
        filepath = filedialog.askopenfilename(filetypes=[("Word Documents", "*.docx")])
        if filepath:
            self.input_file = filepath
            self.dnd_lbl.configure(text=f"Selected: {os.path.basename(filepath)}", text_color=TEXT_MAIN)
            if not self.output_file:
                dir_name = os.path.dirname(filepath)
                base_name = os.path.basename(filepath)
                self.output_file = os.path.join(dir_name, f"formatted_{base_name}")
                self.output_lbl.configure(text=os.path.basename(self.output_file), text_color=TEXT_MAIN)

    def select_output_folder(self):
        dirpath = filedialog.askdirectory(title="Select Save Folder")
        if dirpath:
            base_name = os.path.basename(self.input_file) if self.input_file else "document.docx"
            if not base_name.startswith("formatted_"):
                base_name = f"formatted_{base_name}"
            self.output_file = os.path.join(dirpath, base_name)
            self.output_lbl.configure(text=f"Save Folder: {dirpath}", text_color=TEXT_MAIN)

    def select_output_file(self):
        filepath = filedialog.asksaveasfilename(title="Specify Custom Output Filename", defaultextension=".docx", filetypes=[("Word Documents", "*.docx")])
        if filepath:
            self.output_file = filepath
            self.output_lbl.configure(text=f"File: {os.path.basename(filepath)}", text_color=TEXT_MAIN)

    def _on_drop_single(self, event):
        filepath = event.data
        if filepath.startswith('{') and filepath.endswith('}'):
            filepath = filepath[1:-1]
        if filepath.lower().endswith('.docx'):
            self.input_file = filepath
            self.dnd_lbl.configure(text=f"Selected: {os.path.basename(filepath)}", text_color=TEXT_MAIN)
            if not self.output_file:
                dir_name = os.path.dirname(filepath)
                base_name = os.path.basename(filepath)
                self.output_file = os.path.join(dir_name, f"formatted_{base_name}")
                self.output_lbl.configure(text=os.path.basename(self.output_file), text_color=TEXT_MAIN)

    def select_batch_input(self):
        dirpath = filedialog.askdirectory()
        if dirpath:
            self.batch_input_dir = dirpath
            self.batch_in_lbl.configure(text=dirpath, text_color=TEXT_MAIN)

    def select_batch_output(self):
        dirpath = filedialog.askdirectory()
        if dirpath:
            self.batch_output_dir = dirpath
            self.batch_out_lbl.configure(text=dirpath, text_color=TEXT_MAIN)

    def parse_input_for_inspector(self):
        if not self.input_file or not os.path.exists(self.input_file):
            messagebox.showerror("Error", "Please select a valid input document first.")
            return

        try:
            import docx
            self.doc_obj = docx.Document(self.input_file)
            doc_parser = DocumentParser(self.doc_obj)
            elements = list(doc_parser.extract_elements())
            classifier = ParagraphClassifier()
            
            self.classified_elements = []
            for item in self.inspector_tree.get_children():
                self.inspector_tree.delete(item)

            for idx, el in enumerate(elements):
                cls_str, conf, review_needed, explanation = classifier.classify_with_explanation(el)
                el['classification'] = cls_str
                el['confidence'] = conf
                el['review_needed'] = review_needed
                el['explanation'] = explanation
                self.classified_elements.append(el)

                snippet = el['text'][:60] + "..." if len(el['text']) > 60 else el['text']
                conf_str = f"{conf*100:.0f}%" + (" ⚠️" if review_needed else "")
                self.inspector_tree.insert("", "end", iid=str(idx), values=(idx+1, cls_str, conf_str, snippet, explanation))

            self.tables_count = len(doc_parser.extract_tables())
            self.images_count = doc_parser.count_images()
            self.analytics_data = DocumentAnalytics.analyze(self.doc_obj, self.classified_elements, self.tables_count, self.images_count)

            # Update Stat Cards
            self.card_words.lbl_val.configure(text=str(self.analytics_data["total_words"]))
            self.card_readability.lbl_val.configure(text=str(self.analytics_data["flesch_score"]))
            self.card_grade.lbl_val.configure(text=str(self.analytics_data["fk_grade"]))
            self.card_time.lbl_val.configure(text=f"{self.analytics_data['reading_time_mins']} m")

            # Update Analytics Textbox
            self.analytics_text.delete("1.0", "end")
            an_str = f"Readability: {self.analytics_data['readability_label']}\n"
            an_str += f"Paragraphs: {self.analytics_data['total_paragraphs']} | Words: {self.analytics_data['total_words']} | Sentences: {self.analytics_data['total_sentences']}\n"
            an_str += f"Tables: {self.tables_count} | Images: {self.images_count}\n\n"
            an_str += "--- Structural Hierarchy ---\n"
            for h in self.analytics_data["heading_hierarchy"]:
                an_str += f"• [{h['level']}] {h['text']}\n"
            self.analytics_text.insert("1.0", an_str)

        except Exception as e:
            messagebox.showerror("Error", f"Inspector error: {e}")

    def apply_manual_override(self):
        selected = self.inspector_tree.selection()
        if not selected:
            messagebox.showwarning("Warning", "Please select a paragraph item from the inspector tree.")
            return

        new_cls = self.override_var.get()
        for iid in selected:
            idx = int(iid)
            if idx < len(self.classified_elements):
                self.classified_elements[idx]['classification'] = new_cls
                self.classified_elements[idx]['confidence'] = 1.0
                self.classified_elements[idx]['review_needed'] = False
                self.classified_elements[idx]['explanation'] = "Manually overridden by user"
                
                snippet = self.classified_elements[idx]['text'][:60]
                self.inspector_tree.item(iid, values=(idx+1, new_cls, "100% (Manual)", snippet, "Manually overridden by user"))

        messagebox.showinfo("Success", f"Overrode {len(selected)} item(s) to '{new_cls}'")

    def start_single_formatting(self):
        if not self.input_file or not os.path.exists(self.input_file):
            messagebox.showerror("Error", "Please select a valid input .docx file.")
            return
        if not self.output_file:
            messagebox.showerror("Error", "Please specify a destination file path.")
            return

        self.cancel_requested = False
        self.start_btn.configure(state="disabled", text="PROCESSING DOCUMENT...")
        self.single_status_lbl.configure(text="Processing document...", text_color=ACCENT_BLUE)
        self.progress_bar.set(0.05)


        threading.Thread(target=self._run_single_formatting_thread, daemon=True).start()

    def _run_single_formatting_thread(self):
        try:
            start_time = time.time()
            import docx

            self.log_message(f"🚀 Starting Document Formatting: {os.path.basename(self.input_file)}")
            doc = docx.Document(self.input_file)

            if self.cancel_requested: return

            self.after(0, self.single_status_lbl.configure, {"text": "Cleaning typography and formatting..."})
            self.log_message("🧹 Cleaning document typography (em-dashes, quotes, spacing)...")

            cleaner = DocumentCleaner(
                doc,
                title_case_headings=self.opt_title_case.get(),
                bullet_standardization=self.opt_bullet_clean.get(),
                trailing_whitespace=self.opt_trailing_clean.get()
            )
            spacing_errors = cleaner.clean()
            self.log_message(f"✓ Fixed {spacing_errors} typography/redundancy errors.")
            self.after(0, self.progress_bar.set, 0.25)

            if self.cancel_requested: return

            self.after(0, self.single_status_lbl.configure, {"text": "Classifying paragraph structures..."})
            self.log_message("🧠 Classifying elements via ML + NLP rules...")
            doc_parser = DocumentParser(doc)
            elements = list(doc_parser.extract_elements())
            tables = doc_parser.extract_tables()
            images_count = doc_parser.count_images()

            if not self.classified_elements or len(self.classified_elements) != len(elements):
                classifier = ParagraphClassifier()
                self.classified_elements = []
                for el in elements:
                    cls_str, conf, review_needed, exp = classifier.classify_with_explanation(el)
                    el['classification'] = cls_str
                    el['confidence'] = conf
                    el['review_needed'] = review_needed
                    el['explanation'] = exp
                    self.classified_elements.append(el)

            self.log_message(f"✓ Processed {len(elements)} elements. Found {len(tables)} tables, {images_count} images.")
            self.after(0, self.progress_bar.set, 0.55)

            if self.cancel_requested: return

            # Watermark application
            wm_text = self.opt_watermark.get().strip()
            if wm_text:
                from watermark import WatermarkGenerator
                self.log_message(f"💧 Applying Watermark: '{wm_text}'...")
                WatermarkGenerator.apply_watermark(doc, wm_text)

            # Acronyms & Glossary Table Insertion
            if self.opt_acronyms.get():
                from acronym_extractor import AcronymGlossaryGenerator
                acrs = AcronymGlossaryGenerator.extract_acronyms(self.classified_elements if self.classified_elements else elements)
                if acrs:
                    self.log_message(f"📑 Inserting Acronyms Glossary Table ({len(acrs)} terms)...")
                    AcronymGlossaryGenerator.insert_acronyms_table(doc, acrs)

            # Offline LaTeX Math Conversion
            from math_converter import LaTeXMathConverter
            maths_found = LaTeXMathConverter.process_latex_math(doc)
            if maths_found > 0:
                self.log_message(f"📐 Converted {maths_found} LaTeX formulas to OMML Math Equations.")

            # Table Chart Generator
            from table_visualizer import TableStylingEngine
            charts_count = TableStylingEngine.generate_chart_for_numeric_tables(doc)
            if charts_count > 0:
                self.log_message(f"📊 Embedded {charts_count} Data Charts from numeric tables.")

            # Cover Page Generation
            if self.opt_cover_page.get():
                from cover_page import CoverPageGenerator
                self.log_message("📝 Inserting Title & Cover Page...")
                t_str = os.path.splitext(os.path.basename(self.input_file))[0]
                CoverPageGenerator.insert_cover_page(doc, title=t_str, preset_id=self.active_preset_id)


            preset = self.preset_manager.get_preset(self.active_preset_id)
            self.after(0, self.single_status_lbl.configure, {"text": f"Applying '{preset.name}' typography..."})
            self.log_message(f"📐 Applying Preset '{preset.name}' layout & OXML dynamic headers/footers...")

            formatter = DocumentFormatter(doc, preset=preset)
            formatter.format_elements(self.classified_elements)
            formatter.save(self.output_file)
            self.log_message(f"✓ Saved formatted DOCX: {os.path.basename(self.output_file)}")

            if self.cancel_requested: return

            base_out, _ = os.path.splitext(self.output_file)
            analytics_data = DocumentAnalytics.analyze(doc, self.classified_elements, len(tables), images_count)
            exporter = DocumentExporter(self.classified_elements, title=os.path.splitext(os.path.basename(self.input_file))[0])

            # Side-by-Side Diff
            if self.opt_generate_diff.get():
                from diff_tool import DocumentDiffTool
                diff_path = base_out + "_diff.html"
                DocumentDiffTool.generate_html_diff(elements, self.classified_elements, diff_path)
                self.log_message(f"  - Generated Side-by-Side Diff: {os.path.basename(diff_path)}")

            # Multi-format Exporter
            opt_exp = self.opt_export_format.get()
            if opt_exp != "DOCX Only":
                self.log_message("📤 Generating requested export formats...")
                if "EPUB" in opt_exp or "All Formats" in opt_exp:
                    exporter.export_epub(base_out + ".epub")
                    self.log_message(f"  - EPUB eBook generated: {os.path.basename(base_out)}.epub")
                if "HTML Audit" in opt_exp or "All Formats" in opt_exp:
                    exporter.export_audit_report(base_out + "_audit.html", analytics_data)
                    self.log_message(f"  - HTML Audit Report generated: {os.path.basename(base_out)}_audit.html")
                if "All Formats" in opt_exp:
                    exporter.export_markdown(base_out + ".md")
                    exporter.export_txt(base_out + ".txt")
                    exporter.export_html(base_out + ".html")
                    exporter.export_pdf_fallback(self.output_file, base_out + ".pdf")
                    self.log_message("  - Markdown, Text, HTML, and PDF exports completed.")

            # 1-Click Zip Export
            if self.opt_zip_export.get():
                zip_path = base_out + "_bundle.zip"
                exporter.export_zip(zip_path, self.output_file, analytics_data)
                self.log_message(f"  - 📦 Created 1-Click ZIP Archive: {os.path.basename(zip_path)}")

            elapsed = time.time() - start_time
            self.history_manager.add_entry(self.input_file, self.output_file, elapsed, len(elements), preset.name, "Success")
            self.log_message(f"🎉 Job Completed Successfully in {elapsed:.2f}s!")

            self.after(0, self._on_single_success, elapsed, len(elements))


        except Exception as e:
            import traceback
            traceback.print_exc()
            self.log_message(f"❌ Error occurred: {e}")
            self.after(0, self._on_single_error, str(e))

    def _on_single_success(self, elapsed, count):
        self.progress_bar.set(1.0)
        self.single_status_lbl.configure(text=f"Success! Formatted {count} elements in {elapsed:.1f}s", text_color=ACCENT_GREEN)
        self.start_btn.configure(state="normal", text="🚀 START FORMATTING")
        messagebox.showinfo("Success", f"Document formatted successfully in {elapsed:.1f}s!\nSaved to: {self.output_file}")

    def _on_single_error(self, err):
        self.single_status_lbl.configure(text="Error occurred during processing.", text_color="#EF4444")
        self.start_btn.configure(state="normal", text="🚀 START FORMATTING")
        messagebox.showerror("Error", f"An error occurred:\n{err}")

    def start_batch_formatting(self):
        if not self.batch_input_dir or not os.path.exists(self.batch_input_dir):
            messagebox.showerror("Error", "Please select a valid source folder.")
            return
        if not self.batch_output_dir:
            messagebox.showerror("Error", "Please select an output folder.")
            return

        threading.Thread(target=self._run_batch_thread, daemon=True).start()

    def _run_batch_thread(self):
        try:
            docx_files = [f for f in os.listdir(self.batch_input_dir) if f.lower().endswith('.docx')]
            self.after(0, self.batch_status_lbl.configure, {"text": f"Processing {len(docx_files)} files..."})

            pm = PresetManager()
            preset = pm.get_preset(self.active_preset_id)

            for idx, fname in enumerate(docx_files):
                inp = os.path.join(self.batch_input_dir, fname)
                out = os.path.join(self.batch_output_dir, f"formatted_{fname}")

                import docx
                doc = docx.Document(inp)
                cleaner = DocumentCleaner(doc)
                cleaner.clean()
                doc_parser = DocumentParser(doc)
                elements = list(doc_parser.extract_elements())

                classifier = ParagraphClassifier()
                cls_els = []
                for el in elements:
                    cls_str, conf, review_needed = classifier.classify(el)
                    el['classification'] = cls_str
                    el['confidence'] = conf
                    el['review_needed'] = review_needed
                    cls_els.append(el)

                formatter = DocumentFormatter(doc, preset=preset)
                formatter.format_elements(cls_els)
                formatter.save(out)

            self.after(0, self.batch_status_lbl.configure, {"text": f"Successfully batch formatted {len(docx_files)} files!", "text_color": ACCENT_GREEN})
            self.after(0, messagebox.showinfo, "Batch Complete", f"Successfully formatted {len(docx_files)} files.")

        except Exception as e:
            self.after(0, messagebox.showerror, "Batch Error", str(e))

if __name__ == "__main__":
    app = App()
    app.mainloop()
